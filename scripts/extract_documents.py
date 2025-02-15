import os
import pdfplumber
from docx import Document
from llama_index.core import GPTVectorStoreIndex, Document as LI_Document
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core import Settings
def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding="utf-8") as file:
        return file.read()

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_all_text(data_folder):
    all_texts = []
    for folder, _, files in os.walk(data_folder):
        for file in files:
            file_path = os.path.join(folder, file)
            if file.endswith('.txt'):
                all_texts.append(extract_text_from_txt(file_path))
            elif file.endswith('.docx'):
                all_texts.append(extract_text_from_docx(file_path))
            elif file.endswith('.pdf'):
                all_texts.append(extract_text_from_pdf(file_path))
    return all_texts

def build_llama_index(data_folder, output_file="llama_index_data.txt"):
    # Extract all text from the data folder
    texts = extract_all_text(data_folder)
    
    # Create LlamaIndex documents from each text chunk
    li_docs = []
    for idx, txt in enumerate(texts):
        # You can further split the text here if needed for better retrieval
        li_docs.append(LI_Document(text=txt, doc_id=f"doc_{idx}"))
    
    # Set up a local embedding model using HuggingFaceEmbeddings.
    # "all-MiniLM-L6-v2" is a popular, lightweight embedding model that runs locally.
    embed_model = HuggingFaceEmbedding(model_name="all-MiniLM-L6-v2")
    Settings.embed_model = embed_model
    
    # Build the vector index using the local embeddings
    index = GPTVectorStoreIndex(li_docs)
    
   # Access the document store from the index
    docstore = index.docstore
    
    # Retrieve all documents from the document store
    all_docs = docstore.docs.values()
    # Write the text of each document to the output file
    with open(output_file, "w", encoding="utf-8") as f:
        for doc in all_docs:
            f.write(doc.text + "\n\n")
    print(f"LlamaIndex data exported to {output_file}")

if __name__ == "__main__":
    # Ensure your data folder (with PDFs, DOCX, and TXT files) is at "data/documents"
    data_folder = "data/documents"
    build_llama_index(data_folder)
